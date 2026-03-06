from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC  = r'C:\Users\schmi\Desktop\locla copied\Contempt Appeal\APPENDIX_PAGES\FR_Dkt316-317Comparison.docx'
DEST = r'C:\Users\schmi\Desktop\locla copied\Contempt Appeal\APPENDIX_PAGES\FR_Dkt316-317Comparison.docx'

doc = Document(SRC)

# ── helpers ──────────────────────────────────────────────────────────────────
def clear_para(p):
    for r in list(p._p.findall(qn('w:r'))):
        p._p.remove(r)

def set_para_fmt(p, sb=0, sa=4, ls=None, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(sb)
    pf.space_after  = Pt(sa)
    if ls:
        pf.line_spacing = Pt(ls)
    if align is not None:
        p.alignment = align

def add_run(p, text, size, bold=False, color=None):
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    else:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return run

def set_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'),  'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)

def set_padding(cell, t=70, b=70, l=110, r=110):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for side, val in [('top', t), ('bottom', b), ('left', l), ('right', r)]:
        e = tcMar.find(qn('w:' + side))
        if e is None:
            e = OxmlElement('w:' + side)
            tcMar.append(e)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')

def set_col_width(cell, w):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(w))
    tcW.set(qn('w:type'), 'dxa')

def remove_fixed_height(row):
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is not None:
        h = trPr.find(qn('w:trHeight'))
        if h is not None:
            trPr.remove(h)

# ── 1. BODY PARAGRAPHS ───────────────────────────────────────────────────────
p0 = doc.paragraphs[0]
clear_para(p0)
set_para_fmt(p0, sb=0, sa=3, align=1)
add_run(p0, 'COMPARISON: PROPOSED ORDER (Dkt. 316) vs. FINAL ORDER (Dkt. 317)', 13, bold=True)

p1 = doc.paragraphs[1]
clear_para(p1)
set_para_fmt(p1, sb=0, sa=2, align=1)
add_run(p1, 'Four Material Modifications Added Sua Sponte \u2014 Without Notice to Appellant', 10.5)

p2 = doc.paragraphs[2]
clear_para(p2)
set_para_fmt(p2, sb=0, sa=5, align=1)
add_run(p2, 'Red text marks language added in Dkt.\u00a0317 not present in the proposed order (Dkt.\u00a0316).',
        9.5, color=(0x55, 0x55, 0x55))

p6 = doc.paragraphs[6]
clear_para(p6)
set_para_fmt(p6, sb=5, sa=2, ls=12)
add_run(p6, 'Summary: ', 9.5, bold=True)
add_run(p6, ('Every modification made the order harsher; none were requested by Respondent. '
             'The 7-day hold under Local Rule\u00a0402(11) was not observed, denying Appellant '
             'any opportunity to review or object before entry.'), 9.5)

p7 = doc.paragraphs[7]
clear_para(p7)
set_para_fmt(p7, sb=1, sa=0, ls=12)
add_run(p7, ('All four modifications were imposed sua sponte, without notice or hearing, '
             'violating procedural due process.'), 9.5, bold=True)

# ── 2. TABLE DATA ─────────────────────────────────────────────────────────────
BLACK = (0x1A, 0x1A, 0x1A)
WHITE = (0xFF, 0xFF, 0xFF)
RED   = (0xC0, 0x00, 0x00)
GRAY  = (0x88, 0x88, 0x88)

# Each row: list of 3 cells; each cell: list of (text, bold, color, linebreak_after)
TABLE_DATA = [
    # Row 0 — header
    [
        [('Dkt.\u00a0316 \u2014 Proposed Order', True, WHITE, True),
         ('Filed Jan.\u00a07, 2026 by Atty. Reniero', False, WHITE, False)],
        [('Dkt.\u00a0317 \u2014 Final Order', True, WHITE, True),
         ('Signed Jan.\u00a07 (Hanson); Jan.\u00a08 (Hilton)', False, WHITE, False)],
        [('Modification', True, WHITE, False)],
    ],
    # Row 1
    [
        [('\u00a7\u20093: \u201c\u2026jail ', False, BLACK, False),
         ('with the right to purge', True, BLACK, False),
         (' as follows:\u201d', False, BLACK, False)],
        [('\u00a7\u20093: \u201c\u2026jail ', False, BLACK, False),
         ('without Huber privileges, consecutive to any other sentence, ', True, RED, False),
         ('with the right to purge as follows:\u201d', False, BLACK, False)],
        [('ADDED: ', True, RED, False),
         ('\u201cwithout Huber privileges\u201d \u2014 departs LR\u00a0211 default.', False, BLACK, True),
         ('ADDED: ', True, RED, False),
         ('\u201cconsecutive to any other sentence\u201d \u2014 not requested.', False, BLACK, False)],
    ],
    # Row 2
    [
        [('[No corresponding provision]', False, GRAY, False)],
        [('\u00a7\u20093(b): ', True, RED, False),
         ('\u201cIf Mr. Alexander does not pay, Ms. Anderson may submit a sworn statement\u2026 '
          'If he does not reply, ', False, RED, False),
         ('a bench warrant will be issued.', True, RED, False),
         ('\u201d', False, RED, False)],
        [('ENTIRE PARAGRAPH ADDED. ', True, RED, False),
         ('Self-executing bench warrant \u2014 not proposed by Respondent, '
          'not discussed at any hearing.', False, BLACK, False)],
    ],
    # Row 3
    [
        [('\u00a7\u20094: \u201cA judicial lien shall be placed on any proceeds '
          'in Lennon v.\u00a0ALKAR-Rapidpak\u2026\u201d', False, BLACK, False)],
        [('\u00a7\u20094: \u201cA judicial lien ', False, BLACK, False),
         ('in this amount, and any subsequent awards of attorney fees imposed in '
          'connection with enforcement or defense of this order, ', True, RED, False),
         ('shall be placed on any proceeds\u2026\u201d', False, BLACK, False)],
        [('EXPANDED: ', True, RED, False),
         ('lien extended to future attorney fees \u2014 not in proposed order.',
          False, BLACK, False)],
    ],
    # Row 4
    [
        [('Drafted by Atty. Reniero; no judicial signature block.', False, BLACK, False)],
        [('Signed: Commr. Hanson (Jan.\u00a07); Judge Hilton (Jan.\u00a08).', False, RED, False)],
        [('One-day turnaround bypassed LR\u00a0402(11) 7-day hold.', False, BLACK, False)],
    ],
]

ROW_FILLS  = ['2B3D4F', 'FFFFFF', 'FEF3F0', 'FFFFFF', 'FEF3F0']
COL_WIDTHS = [3375, 3375, 2250]

tbl = doc.tables[0]

for ri, row in enumerate(tbl.rows):
    remove_fixed_height(row)
    fill    = ROW_FILLS[ri]
    is_hdr  = (ri == 0)
    sz      = 9.0 if is_hdr else 9.5

    for ci, cell in enumerate(row.cells):
        set_shading(cell, fill)
        set_padding(cell, t=65, b=65, l=100, r=100)
        set_col_width(cell, COL_WIDTHS[ci])

        # Clear
        for p in cell.paragraphs:
            clear_para(p)

        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        pf.line_spacing = Pt(11.5) if is_hdr else Pt(12.5)

        runs_data = TABLE_DATA[ri][ci]
        for (text, bold, color, lb) in runs_data:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(sz)
            run.bold = bold
            run.font.color.rgb = RGBColor(*color)
            if lb:
                br = OxmlElement('w:br')
                run._r.append(br)

# ── 3. TABLE PROPERTIES ──────────────────────────────────────────────────────
tblPr = tbl._tbl.find(qn('w:tblPr'))

tblW = tblPr.find(qn('w:tblW'))
if tblW is None:
    tblW = OxmlElement('w:tblW')
    tblPr.append(tblW)
tblW.set(qn('w:w'), '9000')
tblW.set(qn('w:type'), 'dxa')

tblInd = tblPr.find(qn('w:tblInd'))
if tblInd is not None:
    tblInd.set(qn('w:w'), '0')

tblGrid = tbl._tbl.find(qn('w:tblGrid'))
if tblGrid is not None:
    for i, gc in enumerate(tblGrid.findall(qn('w:gridCol'))):
        if i < len(COL_WIDTHS):
            gc.set(qn('w:w'), str(COL_WIDTHS[i]))

tblBdr = tblPr.find(qn('w:tblBorders'))
if tblBdr is None:
    tblBdr = OxmlElement('w:tblBorders')
    tblPr.append(tblBdr)
for side in ['top', 'left', 'bottom', 'right']:
    e = tblBdr.find(qn('w:' + side))
    if e is None:
        e = OxmlElement('w:' + side)
        tblBdr.append(e)
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '12')
    e.set(qn('w:color'), '1a1a1a')
for side in ['insideH', 'insideV']:
    e = tblBdr.find(qn('w:' + side))
    if e is None:
        e = OxmlElement('w:' + side)
        tblBdr.append(e)
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '4')
    e.set(qn('w:color'), 'AAAAAA')

# ── 4. MARGINS ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Inches(0.9)
sec.bottom_margin = Inches(0.9)
sec.left_margin   = Inches(1.1)
sec.right_margin  = Inches(1.0)

# ── 5. SAVE ───────────────────────────────────────────────────────────────────
doc.save(DEST)
print('Done:', DEST)
